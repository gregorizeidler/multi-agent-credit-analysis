"""Processador de documentos financeiros com OCR e extração de texto."""

import io
import os
import tempfile
from typing import Dict, List, Optional, Tuple

import fitz  # PyMuPDF
import pytesseract
from docx import Document
from loguru import logger
from PIL import Image

from ..models.schemas import DocumentType


class DocumentProcessor:
    """Processador de documentos com suporte a PDF, DOCX e imagens."""
    
    def __init__(self):
        # Configurar Tesseract para português
        self.tesseract_config = '--oem 3 --psm 6 -l por'
        
        # Palavras-chave para identificar tipos de documentos
        self.document_keywords = {
            DocumentType.BALANCE_SHEET: [
                'balanço patrimonial', 'ativo', 'passivo', 'patrimônio líquido',
                'ativo circulante', 'passivo circulante', 'imobilizado'
            ],
            DocumentType.INCOME_STATEMENT: [
                'demonstração do resultado', 'dre', 'receita líquida', 'lucro líquido',
                'receita bruta', 'custos', 'despesas operacionais', 'ebitda'
            ],
            DocumentType.CASH_FLOW: [
                'fluxo de caixa', 'demonstração dos fluxos de caixa',
                'atividades operacionais', 'atividades de investimento', 'atividades de financiamento'
            ]
        }
    
    async def process_document(self, file_content: bytes, filename: str) -> Tuple[str, DocumentType]:
        """
        Processa um documento e extrai o texto.
        
        Returns:
            Tuple com (texto_extraído, tipo_documento)
        """
        file_extension = self._get_file_extension(filename)
        
        try:
            if file_extension == '.pdf':
                text = await self._extract_from_pdf(file_content)
            elif file_extension == '.docx':
                text = await self._extract_from_docx(file_content)
            elif file_extension in ['.png', '.jpg', '.jpeg', '.tiff']:
                text = await self._extract_from_image(file_content)
            else:
                raise ValueError(f"Formato de arquivo não suportado: {file_extension}")
            
            # Identifica o tipo do documento
            doc_type = self._identify_document_type(text)
            
            logger.info(f"Documento processado: {filename} -> {doc_type}")
            return text, doc_type
            
        except Exception as e:
            logger.error(f"Erro ao processar documento {filename}: {e}")
            raise
    
    def _get_file_extension(self, filename: str) -> str:
        """Extrai a extensão do arquivo."""
        return os.path.splitext(filename.lower())[1]
    
    async def _extract_from_pdf(self, file_content: bytes) -> str:
        """Extrai texto de um PDF usando PyMuPDF e OCR como fallback."""
        try:
            # Primeira tentativa: extrair texto diretamente
            doc = fitz.open(stream=file_content, filetype="pdf")
            text_parts = []
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text = page.get_text()
                
                # Se não há texto, usar OCR
                if not text.strip():
                    # Converter página para imagem
                    pix = page.get_pixmap()
                    img_data = pix.tobytes("png")
                    
                    # OCR na imagem
                    image = Image.open(io.BytesIO(img_data))
                    ocr_text = pytesseract.image_to_string(image, config=self.tesseract_config)
                    text_parts.append(ocr_text)
                else:
                    text_parts.append(text)
            
            doc.close()
            return '\n\n'.join(text_parts)
            
        except Exception as e:
            logger.error(f"Erro ao processar PDF: {e}")
            raise
    
    async def _extract_from_docx(self, file_content: bytes) -> str:
        """Extrai texto de um documento DOCX."""
        try:
            doc = Document(io.BytesIO(file_content))
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Processar tabelas também
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(' | '.join(row_text))
            
            return '\n\n'.join(text_parts)
            
        except Exception as e:
            logger.error(f"Erro ao processar DOCX: {e}")
            raise
    
    async def _extract_from_image(self, file_content: bytes) -> str:
        """Extrai texto de uma imagem usando OCR."""
        try:
            image = Image.open(io.BytesIO(file_content))
            
            # Pré-processamento da imagem para melhorar OCR
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # OCR
            text = pytesseract.image_to_string(image, config=self.tesseract_config)
            
            return text
            
        except Exception as e:
            logger.error(f"Erro ao processar imagem: {e}")
            raise
    
    def _identify_document_type(self, text: str) -> DocumentType:
        """Identifica o tipo de documento baseado no conteúdo."""
        text_lower = text.lower()
        
        # Conta palavras-chave para cada tipo
        type_scores = {}
        
        for doc_type, keywords in self.document_keywords.items():
            score = sum(1 for keyword in keywords if keyword in text_lower)
            type_scores[doc_type] = score
        
        # Retorna o tipo com maior score
        if type_scores:
            best_type = max(type_scores, key=type_scores.get)
            if type_scores[best_type] > 0:
                return best_type
        
        return DocumentType.OTHER
    
    def extract_financial_data(self, text: str, doc_type: DocumentType) -> Dict[str, Optional[float]]:
        """
        Extrai dados financeiros específicos baseado no tipo de documento.
        
        Esta é uma versão simplificada. Em produção, usaria ML/NLP mais sofisticado.
        """
        financial_data = {}
        
        if doc_type == DocumentType.BALANCE_SHEET:
            financial_data.update(self._extract_balance_sheet_data(text))
        elif doc_type == DocumentType.INCOME_STATEMENT:
            financial_data.update(self._extract_income_statement_data(text))
        elif doc_type == DocumentType.CASH_FLOW:
            financial_data.update(self._extract_cash_flow_data(text))
        
        return financial_data
    
    def _extract_balance_sheet_data(self, text: str) -> Dict[str, Optional[float]]:
        """Extrai dados do balanço patrimonial."""
        import re
        
        patterns = {
            'total_assets': r'ativo\s+total\s*[:\-]?\s*([\d\.,]+)',
            'current_assets': r'ativo\s+circulante\s*[:\-]?\s*([\d\.,]+)',
            'total_liabilities': r'passivo\s+total\s*[:\-]?\s*([\d\.,]+)',
            'current_liabilities': r'passivo\s+circulante\s*[:\-]?\s*([\d\.,]+)',
            'equity': r'patrimônio\s+líquido\s*[:\-]?\s*([\d\.,]+)',
        }
        
        return self._extract_values_with_patterns(text, patterns)
    
    def _extract_income_statement_data(self, text: str) -> Dict[str, Optional[float]]:
        """Extrai dados da DRE."""
        patterns = {
            'revenue': r'receita\s+(?:líquida|total)\s*[:\-]?\s*([\d\.,]+)',
            'gross_profit': r'lucro\s+bruto\s*[:\-]?\s*([\d\.,]+)',
            'operating_profit': r'lucro\s+operacional\s*[:\-]?\s*([\d\.,]+)',
            'net_profit': r'lucro\s+líquido\s*[:\-]?\s*([\d\.,]+)',
        }
        
        return self._extract_values_with_patterns(text, patterns)
    
    def _extract_cash_flow_data(self, text: str) -> Dict[str, Optional[float]]:
        """Extrai dados do fluxo de caixa."""
        patterns = {
            'operating_cash_flow': r'(?:fluxo\s+)?(?:de\s+)?caixa\s+(?:das\s+)?atividades\s+operacionais\s*[:\-]?\s*([\d\.,]+)',
            'investing_cash_flow': r'(?:fluxo\s+)?(?:de\s+)?caixa\s+(?:das\s+)?atividades\s+de\s+investimento\s*[:\-]?\s*([\d\.,]+)',
            'financing_cash_flow': r'(?:fluxo\s+)?(?:de\s+)?caixa\s+(?:das\s+)?atividades\s+de\s+financiamento\s*[:\-]?\s*([\d\.,]+)',
        }
        
        return self._extract_values_with_patterns(text, patterns)
    
    def _extract_values_with_patterns(self, text: str, patterns: Dict[str, str]) -> Dict[str, Optional[float]]:
        """Extrai valores numéricos usando regex patterns."""
        import re
        
        results = {}
        text_lower = text.lower()
        
        for key, pattern in patterns.items():
            matches = re.findall(pattern, text_lower, re.IGNORECASE)
            
            if matches:
                try:
                    # Limpa e converte o valor
                    value_str = matches[0].replace('.', '').replace(',', '.')
                    value = float(value_str)
                    results[key] = value
                except (ValueError, IndexError):
                    results[key] = None
            else:
                results[key] = None
        
        return results


# Singleton instance
document_processor = DocumentProcessor()